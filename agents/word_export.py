from docx import Document
from io import BytesIO


def create_word_file(title, data):

    doc = Document()
    doc.add_heading(title, level=1)

    if isinstance(data, dict):
        for k, v in data.items():
            doc.add_heading(str(k), level=2)
            doc.add_paragraph(str(v))

    elif isinstance(data, list):
        for i, item in enumerate(data, 1):
            doc.add_paragraph(f"{i}. {item}")

    elif isinstance(data, str):
        doc.add_paragraph(data)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
from docx import Document
from io import BytesIO


def create_viva_word(viva_data):

    doc = Document()
    doc.add_heading("Viva Questions & Answers", level=1)

    questions = viva_data.get("questions", [])

    for i, q in enumerate(questions, 1):

        doc.add_heading(f"Q{i}: {q.get('question','')}", level=2)

        doc.add_paragraph("Answer:")
        doc.add_paragraph(q.get("answer", ""))

        doc.add_paragraph(f"Difficulty: {q.get('difficulty','')}")
        doc.add_paragraph(f"Category: {q.get('category','')}")

        doc.add_paragraph("-" * 40)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer